"""
doc_builder.py
Membangun file .docx dokumentasi survey secara in-memory.
Tidak ada file yang ditulis ke disk.
"""

import io
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


# ── Konstanta ──────────────────────────────────────────────────────────────────
PAGE_W_CM   = 21.0   # A4 lebar
PAGE_H_CM   = 29.7   # A4 tinggi
MARGIN_CM   = 2.0    # margin semua sisi
CONTENT_W   = PAGE_W_CM - (MARGIN_CM * 2)   # 17 cm

FONT_NAME   = "Calibri (Body)"
FONT_SIZE   = 10.5

# Ukuran sel tabel foto (dari Word, dalam inch)
CELL_H_IN   = 2.67
CELL_W_IN   = 3.25

# Ukuran foto dalam sel (dalam inch), sesuai orientasi
PHOTO_PORTRAIT_W  = Inches(1.91)
PHOTO_PORTRAIT_H  = Inches(2.55)
PHOTO_LANDSCAPE_W = Inches(3.10)
PHOTO_LANDSCAPE_H = Inches(2.32)

PHOTO_PADDING_CM  = 0.2   # padding foto dalam sel (masih dipakai untuk label APD)


# ── Helper: set cell borders ───────────────────────────────────────────────────
def _set_cell_border(cell, **kwargs):
    """Set border pada cell tabel. kwargs: top, bottom, left, right → dict dengan sz, color, val."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        for k, v in attrs.items():
            el.set(qn(f"w:{k}"), v)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _border_single():
    return {"val": "single", "sz": "4", "space": "0", "color": "000000"}


# ── Helper: set cell vertical alignment ───────────────────────────────────────
def _set_vertical_align(cell, align=WD_ALIGN_VERTICAL.CENTER):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)


# ── Helper: merge sel horizontal (colspan=2) ──────────────────────────────────
def _merge_row_cells(row):
    """Merge kedua cell dalam satu row menjadi satu (colspan 2)."""
    row.cells[0].merge(row.cells[1])


# ── Helper: resize foto ────────────────────────────────────────────────────────
def _prepare_image(file_bytes: bytes) -> tuple[io.BytesIO, str]:
    """
    Apply EXIF rotation, deteksi orientasi, resize proporsional.
    Return (BytesIO, orientasi) → orientasi: 'portrait' atau 'landscape'
    """
    from PIL import ImageOps
    img = Image.open(io.BytesIO(file_bytes))
    img = ImageOps.exif_transpose(img)   # apply rotasi EXIF dulu
    img = img.convert("RGB")

    orientasi = "landscape" if img.width > img.height else "portrait"

    # Resize sesuai orientasi, pertahankan aspek rasio
    if orientasi == "landscape":
        target_w = int(PHOTO_LANDSCAPE_W.inches * 96)
        target_h = int(PHOTO_LANDSCAPE_H.inches * 96)
    else:
        target_w = int(PHOTO_PORTRAIT_W.inches * 96)
        target_h = int(PHOTO_PORTRAIT_H.inches * 96)

    img.thumbnail((target_w, target_h), Image.LANCZOS)

    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    buf.seek(0)
    return buf, orientasi


# ── Helper: insert foto ke dalam cell ─────────────────────────────────────────
def _insert_photo(cell, file_bytes: bytes):
    """Masukkan foto ke dalam cell dengan ukuran sesuai orientasi."""
    buf, orientasi = _prepare_image(file_bytes)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run = para.add_run()
    if orientasi == "landscape":
        run.add_picture(buf, width=PHOTO_LANDSCAPE_W, height=PHOTO_LANDSCAPE_H)
    else:
        run.add_picture(buf, width=PHOTO_PORTRAIT_W, height=PHOTO_PORTRAIT_H)
    _set_vertical_align(cell)


# ── Helper: cell kosong (placeholder abu-abu) ─────────────────────────────────
def _empty_cell(cell):
    """Cell tanpa foto — biarkan kosong, hanya set alignment."""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_vertical_align(cell)


# ── Helper: set lebar kolom tabel ─────────────────────────────────────────────
def _set_col_widths(table, widths_in):
    """Set lebar setiap kolom (dalam cm)."""
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if j < len(widths_in):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(int(widths_in[j] * 1440)))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)


# ── Helper: set tinggi baris ──────────────────────────────────────────────────
def _set_row_height(row, height_in):
    """Set tinggi baris dalam inch."""
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_in * 1440)))  # 1440 twips per inch
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


# ── Helper: run teks berformat ────────────────────────────────────────────────
def _add_run(para, text, bold=False, size=FONT_SIZE, color=None):
    run = para.add_run(text)
    run.font.name      = FONT_NAME
    run.font.size      = Pt(size)
    run.font.bold      = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


# ── Bagian: header info ────────────────────────────────────────────────────────
def _build_header(doc, data: dict):
    """Tambahkan judul dan tabel info ke dokumen."""
    # Judul
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(title_para, "DOKUMENTASI SURVEI CAPEX 2026", bold=True, size=10.5)
    title_para.paragraph_format.space_after  = Pt(6)
    title_para.paragraph_format.space_before = Pt(0)

    # Tabel info: 3 kolom (label | titik dua | nilai)
    fields = [
        ("Hari/Tanggal",                    data.get("hari_tanggal", "")),
        ("Nama Kegiatan",                   data.get("nama_kegiatan", "")),
        ("Jenis Kontrak",                   data.get("jenis_kontrak", "")),
        ("Lokasi",                          data.get("lokasi", "")),
        ("Divisi",                          data.get("divisi", "")),
        ("Subdivisi",                       data.get("subdivisi", "")),
        ("PJK",                             data.get("pjk", "")),
        ("Nilai Terkontrak",                data.get("nilai_terkontrak", "")),
        ("Nilai Terkontrak Final (Jika Ada Sesuai Addendum)", data.get("nilai_terkontrak_final", "")),
        ("Vendor",                          data.get("vendor", "")),
        ("Progress Lap-%",                  data.get("progress", "")),
    ]

    tbl = doc.add_table(rows=len(fields), cols=3)
    tbl.style = "Normal Table"

    # Lebar kolom: label=5cm, titik dua=0.5cm, nilai=sisa
    label_w = 1.81
    sep_w   = 0.19
    val_w   = 4.58

    for i, (label, value) in enumerate(fields):
        row = tbl.rows[i]

        # Hapus border tabel info (tampilan bersih seperti template)
        for cell in row.cells:
            _set_cell_border(cell,
                top={"val": "none", "sz": "0", "space": "0", "color": "auto"},
                bottom={"val": "none", "sz": "0", "space": "0", "color": "auto"},
                left={"val": "none", "sz": "0", "space": "0", "color": "auto"},
                right={"val": "none", "sz": "0", "space": "0", "color": "auto"},
            )

        # Kolom 0: label
        p0 = row.cells[0].paragraphs[0]
        _add_run(p0, label, bold=True)
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after  = Pt(0)

        # Kolom 1: titik dua
        p1 = row.cells[1].paragraphs[0]
        _add_run(p1, ":")

        # Kolom 2: nilai
        p2 = row.cells[2].paragraphs[0]
        _add_run(p2, value)

    # Set lebar kolom info
    _set_col_widths(tbl, [label_w, sep_w, val_w])

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Bagian: tabel foto ────────────────────────────────────────────────────────
def _build_photo_table(doc, photos: list, apd_photos: list, include_apd: bool):
    """
    Bangun tabel foto dinamis.

    Layout:
    - Foto biasa disusun 2 per baris
    - Jika ganjil → baris terakhir di-merge (foto di tengah)
    - Jika include_apd → tambah row label "Penggunaan APD" lalu foto APD
    """
    ROW_H_IN = CELL_H_IN

    # Hitung total baris yang dibutuhkan
    n_biasa    = len(photos)
    n_apd      = len(apd_photos) if include_apd else 0
    rows_biasa = (n_biasa + 1) // 2   # ceil division
    rows_apd   = (n_apd + 1) // 2 if include_apd else 0

    total_rows = rows_biasa
    if include_apd:
        total_rows += 1 + rows_apd   # +1 untuk label APD

    tbl = doc.add_table(rows=total_rows, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set lebar kolom foto
    for row in tbl.rows:
        for j, cell in enumerate(row.cells):
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW  = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(CELL_W_IN * 1440)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            _set_cell_border(cell,
                top=_border_single(), bottom=_border_single(),
                left=_border_single(), right=_border_single(),
            )

    # ── Isi foto biasa ─────────────────────────────────────────────────────────
    for row_i in range(rows_biasa):
        row      = tbl.rows[row_i]
        idx_kiri = row_i * 2
        idx_kanan= idx_kiri + 1
        is_last  = (row_i == rows_biasa - 1)
        is_odd   = (n_biasa % 2 == 1)

        _set_row_height(row, ROW_H_IN)

        if is_last and is_odd:
            # Merge 2 kolom, foto di tengah
            _merge_row_cells(row)
            _insert_photo(row.cells[0], photos[idx_kiri])
        else:
            _insert_photo(row.cells[0], photos[idx_kiri])
            if idx_kanan < n_biasa:
                _insert_photo(row.cells[1], photos[idx_kanan])
            else:
                _empty_cell(row.cells[1])

    # ── Label APD & foto APD ───────────────────────────────────────────────────
    if include_apd:
        apd_label_row_i = rows_biasa
        apd_label_row   = tbl.rows[apd_label_row_i]

        # Merge label APD full width
        _merge_row_cells(apd_label_row)
        label_cell = apd_label_row.cells[0]
        _set_row_height(apd_label_row, 0.2)

        # Border label
        _set_cell_border(label_cell,
            top=_border_single(), bottom=_border_single(),
            left=_border_single(), right=_border_single(),
        )

        p_label = label_cell.paragraphs[0]
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p_label, "Penggunaan APD", bold=True)

        # Foto APD
        for apd_i in range(rows_apd):
            row      = tbl.rows[apd_label_row_i + 1 + apd_i]
            idx_kiri = apd_i * 2
            idx_kanan= idx_kiri + 1
            is_last  = (apd_i == rows_apd - 1)
            is_odd   = (n_apd % 2 == 1)

            _set_row_height(row, ROW_H_IN)

            if is_last and is_odd:
                _merge_row_cells(row)
                _insert_photo(row.cells[0], apd_photos[idx_kiri])
            else:
                _insert_photo(row.cells[0], apd_photos[idx_kiri])
                if idx_kanan < n_apd:
                    _insert_photo(row.cells[1], apd_photos[idx_kanan])
                else:
                    _empty_cell(row.cells[1])


# ── Fungsi utama ───────────────────────────────────────────────────────────────
def build_docx(data: dict, photos: list, apd_photos: list, include_apd: bool) -> bytes:
    """
    Generate dokumen Word in-memory dan kembalikan sebagai bytes.

    Args:
        data        : dict berisi field-field header
        photos      : list of bytes, foto biasa (maks 8 jika ada APD, maks 10 jika tidak)
        apd_photos  : list of bytes, foto APD (maks 2)
        include_apd : bool, apakah ada seksi foto APD

    Returns:
        bytes: isi file .docx
    """
    doc = Document()

    # Set margin halaman
    for section in doc.sections:
        section.top_margin    = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin   = Cm(MARGIN_CM)
        section.right_margin  = Cm(MARGIN_CM)

    _build_header(doc, data)
    _build_photo_table(doc, photos, apd_photos, include_apd)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
